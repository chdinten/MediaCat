"""Generate MediaCat technical reference document as a rich Word (.docx) file."""

from __future__ import annotations

import sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ---------------------------------------------------------------------------
# Colour palette — stored as (r, g, b) tuples; converted to RGBColor on use
# ---------------------------------------------------------------------------
C_NAVY    = (0x1A, 0x2E, 0x4A)   # headings / table headers
C_TEAL    = (0x00, 0x7A, 0x87)   # accent / h2
C_AMBER   = (0xD4, 0x7E, 0x00)   # callout / warning
C_SLATE   = (0x4A, 0x5B, 0x6C)   # body text
C_WHITE   = (0xFF, 0xFF, 0xFF)
C_LIGHT   = (0xF2, 0xF6, 0xFA)   # table zebra
C_HEADER  = (0x1A, 0x2E, 0x4A)   # table header bg
C_GREEN   = (0x1A, 0x7A, 0x4A)   # ✓ items
C_RED     = (0x9B, 0x1C, 0x1C)   # open items

def _to_rgb(c: tuple) -> RGBColor:
    return RGBColor(c[0], c[1], c[2])

def _to_hex(c: tuple) -> str:
    return f"{c[0]:02X}{c[1]:02X}{c[2]:02X}"

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _set_cell_bg(cell, colour: tuple):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  _to_hex(colour))
    tcPr.append(shd)


def _set_cell_margins(cell, top=50, bottom=50, left=100, right=100):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("bottom", bottom),
                      ("left", left), ("right", right)):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"),    str(val))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)


def _add_run(para, text: str, bold=False, italic=False, size=10,
             colour: tuple | None = None, mono=False):
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size  = Pt(size)
    if colour:
        run.font.color.rgb = _to_rgb(colour)
    if mono:
        run.font.name = "Consolas"
    return run


def _h1(doc, text: str):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = _to_rgb(C_WHITE)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  _to_hex(C_NAVY))
    pPr.append(shd)
    return p


def _h2(doc, text: str):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = _to_rgb(C_TEAL)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    return p


def _h3(doc, text: str):
    p = doc.add_heading(text, level=3)
    p.runs[0].font.color.rgb = _to_rgb(C_NAVY)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    return p


def _body(doc, text: str):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        run.font.color.rgb = _to_rgb(C_SLATE)
        run.font.size      = Pt(10)
    return p


def _bullet(doc, text: str, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent   = Pt(18 * (level + 1))
    p.paragraph_format.space_after   = Pt(2)
    for run in p.runs:
        run.font.size      = Pt(10)
        run.font.color.rgb = _to_rgb(C_SLATE)
    return p


def _code(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.8)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  _to_hex(C_LIGHT))
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name  = "Consolas"
    run.font.size  = Pt(8.5)
    run.font.color.rgb = _to_rgb(C_NAVY)
    return p


def _callout(doc, label: str, text: str, colour: tuple = C_AMBER):
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    w0, w1 = tbl.columns[0], tbl.columns[1]
    w0.width = Cm(0.5)
    w1.width = Cm(14.5)
    c0, c1 = tbl.rows[0].cells
    _set_cell_bg(c0, colour)
    bg1 = (0xFF, 0xFB, 0xF0) if colour == C_AMBER else C_LIGHT
    _set_cell_bg(c1, bg1)
    p = c1.paragraphs[0]
    _add_run(p, f"{label}  ", bold=True, colour=colour, size=10)
    _add_run(p, text, size=10, colour=C_SLATE)
    doc.add_paragraph()


def _simple_table(doc, headers: list[str], rows: list[list[str]],
                  col_widths: list[float] | None = None):
    n = len(headers)
    tbl = doc.add_table(rows=1 + len(rows), cols=n)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in tbl.columns[i].cells:
                cell.width = Cm(w)
    # Header row
    hrow = tbl.rows[0]
    for i, hdr in enumerate(headers):
        cell = hrow.cells[i]
        _set_cell_bg(cell, C_HEADER)
        _set_cell_margins(cell)
        p = cell.paragraphs[0]
        run = p.add_run(hdr)
        run.bold = True
        run.font.size      = Pt(9)
        run.font.color.rgb = _to_rgb(C_WHITE)
    # Data rows
    for ri, row_data in enumerate(rows):
        row = tbl.rows[ri + 1]
        bg  = C_LIGHT if ri % 2 == 0 else C_WHITE
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            _set_cell_bg(cell, bg)
            _set_cell_margins(cell)
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.color.rgb = _to_rgb(C_SLATE)
    doc.add_paragraph()
    return tbl


# ---------------------------------------------------------------------------
# Main document builder
# ---------------------------------------------------------------------------

def build(out_path: str = "MediaCat_Technical_Reference.docx"):
    doc = Document()

    # ── Page layout ────────────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width  = Cm(21.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # ── Cover page ─────────────────────────────────────────────────────────
    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover.paragraph_format.space_before = Pt(80)

    run = cover.add_run("MediaCat")
    run.bold = True
    run.font.size      = Pt(42)
    run.font.color.rgb = _to_rgb(C_NAVY)

    doc.add_paragraph()

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(sub, "Technical Reference & User Guide", bold=False, size=20, colour=C_TEAL)

    doc.add_paragraph()

    ver = doc.add_paragraph()
    ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(ver, "Version 0.1  ·  April 2026  ·  CONFIDENTIAL", size=11, colour=C_SLATE)

    doc.add_paragraph()
    doc.add_paragraph()

    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(tagline,
        "Cataloguing platform for vinyl records and compact discs.\n"
        "Python 3.12 · FastAPI · PostgreSQL 16 · MinIO · Redis · OPA",
        size=11, italic=True, colour=C_SLATE)

    doc.add_page_break()

    # ── Table of Contents placeholder ──────────────────────────────────────
    _h1(doc, "Table of Contents")
    toc_items = [
        ("1", "Project Overview & Architecture"),
        ("2", "Installation — Windows WSL2"),
        ("3", "Database Schema"),
        ("4", "Vision & OCR Pipeline"),
        ("5", "Symbol Registry"),
        ("6", "Ingestion Connectors"),
        ("7", "Review Queue"),
        ("8", "Authentication & Security"),
        ("9", "Object Storage (MinIO)"),
        ("10", "Configuration Reference"),
        ("11", "API Endpoints"),
        ("12", "Alembic Migrations"),
        ("13", "Make Targets"),
        ("14", "Open Items & Roadmap"),
        ("15", "Usage Guide"),
    ]
    tbl = doc.add_table(rows=len(toc_items), cols=2)
    tbl.style = "Table Grid"
    for i, (num, title) in enumerate(toc_items):
        r = tbl.rows[i]
        bg = C_LIGHT if i % 2 == 0 else C_WHITE
        c0, c1 = r.cells
        c0.width = Cm(1.5)
        c1.width = Cm(13.5)
        _set_cell_bg(c0, C_NAVY)
        _set_cell_bg(c1, bg)
        _set_cell_margins(c0)
        _set_cell_margins(c1)
        p0 = c0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p0.add_run(num)
        run.bold = True
        run.font.color.rgb = _to_rgb(C_WHITE)
        run.font.size = Pt(10)
        p1 = c1.paragraphs[0]
        run = p1.add_run(title)
        run.font.size = Pt(10)
        run.font.color.rgb = _to_rgb(C_SLATE)
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # 1. PROJECT OVERVIEW & ARCHITECTURE
    # ═══════════════════════════════════════════════════════════════════════
    _h1(doc, "1 · Project Overview & Architecture")

    _h2(doc, "1.1 What Is MediaCat?")
    _body(doc,
        "MediaCat is a hybrid-AI cataloguing platform for physical music media — "
        "vinyl records and compact discs.  It ingests metadata and cover/runout "
        "photography from external sources (Discogs, MusicBrainz, CoverArtArchive), "
        "enriches records with OCR and vision-model transcription, and stores "
        "everything in an append-only token-object registry with mandatory human "
        "review gates.  No AI decision is ever applied automatically: LLMs propose, "
        "humans approve.")

    _h2(doc, "1.2 Technology Stack")
    _simple_table(doc,
        ["Layer", "Technology", "Version / Notes"],
        [
            ["Language",        "Python",                     "3.12, async/await throughout"],
            ["Web Framework",   "FastAPI + Jinja2 + HTMX",   "Server-side rendering, no heavy JS framework"],
            ["Database",        "PostgreSQL",                 "16, pg_trgm + uuid-ossp + btree_gist"],
            ["ORM / Migrations","SQLAlchemy 2 + Alembic",    "Async ORM, deterministic migration names"],
            ["Object Storage",  "MinIO",                     "S3-compatible, content-hash dedup"],
            ["Job Queue",       "Redis 7",                   "BLMOVE atomic dequeue, stale-job reaper"],
            ["Rule Engine",     "Open Policy Agent (OPA)",   "Rego policies + Python fallback"],
            ["Vision / LLM",    "Ollama (local-first)",      "LLaVA / Qwen2-VL; Anthropic API fallback"],
            ["Reverse Proxy",   "Caddy",                     "Automatic TLS (Let's Encrypt)"],
            ["Deployment",      "Docker Compose on WSL2",    "8-service stack; host-mounted data volumes"],
            ["Auth / Security", "Argon2id + TOTP + OPA",     "MFA scaffold; CSP; X-CSRF-Token header"],
        ],
        col_widths=[3.5, 4.0, 7.5],
    )

    _h2(doc, "1.3 Architectural Principles")
    for item in [
        ("Advisory-only AI",        "LLMs and vision models propose updates; humans always confirm."),
        ("Append-only revisions",   "Every change creates a new TokenRevision row; nothing is overwritten."),
        ("Content-addressed storage","Images are keyed by SHA-256 hash in MinIO — duplicates are free."),
        ("Data / code separation",  "All persistent state lives on host-mounted volumes; containers are ephemeral."),
        ("Least-privilege DB roles","The app role cannot ALTER, DROP, or write the audit log."),
        ("Hybrid AI resilience",    "Local Ollama handles 99% of traffic; API fallback fires automatically on failure."),
    ]:
        p = doc.add_paragraph()
        _add_run(p, item[0] + " — ", bold=True, colour=C_NAVY, size=10)
        _add_run(p, item[1], colour=C_SLATE, size=10)
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(3)

    _h2(doc, "1.4 Repository Layout")
    layout = [
        ("mediacat/src/mediacat/",   "Python package root"),
        ("  db/",                    "ORM models, enums, base, engine, symbol helpers"),
        ("  web/",                   "FastAPI app, routes, auth, middleware, templates"),
        ("  vision/",                "VLM adapter, task prompts, candidate matcher"),
        ("  llm/",                   "LLM adapter, tasks, safety, Ollama & Anthropic backends"),
        ("  ingestion/",             "Connector base, Discogs, MusicBrainz, Redis queue, drift"),
        ("  rules/",                 "OPA adapter, local Python fallback"),
        ("  storage/",               "MinIO wrapper, image pipeline, OCR, translation"),
        ("mediacat/alembic/versions/","Sequential migrations (0001 initial, 0002 symbols)"),
        ("mediacat/deploy/",         "Docker Compose, Caddyfile, OPA bundles, backup scripts"),
        ("mediacat/config/",         "app.yaml + connectors.yaml (host-mounted)"),
        ("mediacat/docs/",           "ADRs, section docs, due-diligence report"),
        ("mediacat/tests/",          "Pytest suite (models, vision, ingestion, web, storage…)"),
    ]
    for path, desc in layout:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.3)
        p.paragraph_format.space_after = Pt(1)
        _add_run(p, path, mono=True, size=9, colour=C_NAVY)
        _add_run(p, f"  —  {desc}", size=9, colour=C_SLATE)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # 2. INSTALLATION
    # ═══════════════════════════════════════════════════════════════════════
    _h1(doc, "2 · Installation — Windows WSL2")

    _h2(doc, "2.1 Prerequisites")
    _simple_table(doc,
        ["Requirement", "Minimum Version", "Notes"],
        [
            ["Windows 11 (or 10 21H2+)", "–",       "WSL2 kernel included"],
            ["WSL2 — Ubuntu 24.04 LTS",  "24.04",   "`wsl --install -d Ubuntu-24.04`"],
            ["Docker Desktop for Windows","4.28+",   "Enable WSL2 backend; allocate ≥ 8 GB RAM"],
            ["Git",                       "2.40+",   "Pre-installed in Ubuntu 24.04"],
            ["GNU Make",                  "4.3+",    "`sudo apt-get install make`"],
            ["uv (Python package manager)","0.4+",   "Installed automatically by `make setup`"],
        ],
        col_widths=[4.5, 3.0, 7.5],
    )

    _h2(doc, "2.2 Step-by-Step Installation")

    steps = [
        ("Step 1 — Enable WSL2 & Install Ubuntu",
         'Open PowerShell as Administrator and run:\n'
         'wsl --install -d Ubuntu-24.04\n'
         'Restart when prompted, then set a UNIX username and password.'),
        ("Step 2 — Install Docker Desktop",
         'Download Docker Desktop from docker.com/products/docker-desktop.\n'
         'In Settings → Resources → WSL Integration, enable your Ubuntu distro.'),
        ("Step 3 — Clone the Repository",
         'Inside the Ubuntu terminal:\n'
         'mkdir -p ~/projects && cd ~/projects\n'
         'git clone <your-repo-url> sounddb\n'
         'cd sounddb/mediacat'),
        ("Step 4 — Bootstrap the System",
         'Run the Ubuntu bootstrap script (installs build tools, Tesseract, etc.):\n'
         'make bootstrap\n'
         'This requires sudo — enter your password when prompted.'),
        ("Step 5 — Create the Data Directory Tree",
         'Initialise the persistent data layout on the host:\n'
         'make data-init\n'
         'Default root: ~/data/mediacat  (override with MEDIACAT_DATA_ROOT)'),
        ("Step 6 — Configure Secrets",
         'Copy the example environment file and edit it:\n'
         'cp .env.example .env\n'
         'nano .env\n'
         'Create secret files in ~/data/mediacat/secrets/:\n'
         '  postgres_app_password, minio_root_password,\n'
         '  redis_password, session_secret\n'
         '(Each file should contain only the secret value, no newline.)'),
        ("Step 7 — Set Up the Python Environment",
         'Install all Python dependencies and pre-commit hooks:\n'
         'make setup\n'
         'This uses uv to create a virtualenv at .venv/ and installs\n'
         'all packages from pyproject.toml including dev extras.'),
        ("Step 8 — Start the Stack",
         'Bring up all 8 Docker services:\n'
         'make up\n'
         'Services: Caddy, PostgreSQL 16, MinIO, Redis 7, OPA, app, worker, backup.'),
        ("Step 9 — Run Migrations",
         'Apply the database schema (runs inside the app container):\n'
         'docker compose exec app alembic upgrade head\n'
         'This creates all tables, enums, indexes, and seeds symbol data.'),
        ("Step 10 — Verify the Installation",
         'Open a browser at http://localhost (or https:// if TLS configured).\n'
         'Log in with the dev admin credentials set in .env.\n'
         'Check health endpoints: curl http://localhost/healthz'),
    ]

    for title, body in steps:
        _h3(doc, title)
        lines = body.split("\n")
        for i, line in enumerate(lines):
            if i == 0:
                _body(doc, line)
            else:
                stripped = line.strip()
                if stripped and (stripped.startswith("make") or stripped.startswith("wsl")
                        or stripped.startswith("git") or stripped.startswith("cp")
                        or stripped.startswith("nano") or stripped.startswith("mkdir")
                        or stripped.startswith("docker") or stripped.startswith("curl")
                        or stripped.startswith("cd") or stripped.startswith("sudo")):
                    _code(doc, stripped)
                elif stripped:
                    _body(doc, stripped)

    _h2(doc, "2.3 Useful Docker Commands")
    _simple_table(doc,
        ["Command", "Description"],
        [
            ["make up",             "Start all services in the background"],
            ["make down",           "Stop and remove containers (data volumes preserved)"],
            ["make restart",        "Restart the stack"],
            ["make logs",           "Tail logs from all services"],
            ["make ps",             "Show container status"],
            ["make config-check",   "Validate docker-compose configuration"],
            ["docker compose exec app bash", "Open a shell inside the app container"],
        ],
        col_widths=[6.0, 9.0],
    )

    _h2(doc, "2.4 Running Tests & Quality Gates")
    for cmd, desc in [
        ("make lint",       "Ruff linter — checks code style and imports"),
        ("make format",     "Ruff formatter — auto-fixes formatting"),
        ("make typecheck",  "Mypy strict — full static type checking"),
        ("make security",   "Bandit + pip-audit — security scan"),
        ("make test",       "Full pytest suite with coverage report"),
        ("make test-fast",  "Pytest excluding slow / integration tests"),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        _add_run(p, cmd, mono=True, bold=True, size=10, colour=C_TEAL)
        _add_run(p, f"  —  {desc}", size=10, colour=C_SLATE)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # 3. DATABASE SCHEMA
    # ═══════════════════════════════════════════════════════════════════════
    _h1(doc, "3 · Database Schema")

    _body(doc,
        "PostgreSQL 16 is the primary data store.  All tables use UUID primary keys "
        "(server-generated via uuid_generate_v4()).  Timestamps are stored in UTC "
        "with timezone.  The schema is managed exclusively via Alembic migrations; "
        "never modify the database schema by hand.")

    _h2(doc, "3.1 Entity Relationship Overview")
    er_desc = [
        "Token — the unit of identity (one per unique physical release)",
        "TokenRevision — append-only history; every change creates a new revision",
        "MediaObject — image stored in MinIO, linked to a token",
        "OcrArtifact — OCR text extracted from a MediaObject",
        "Symbol — canonical graphical runout mark (e.g. EMI triangle)",
        "SymbolVariant — visual variant of a canonical Symbol",
        "TokenSymbol — FK index: which symbols appear in a token's runout",
        "ReviewItem — proposed change awaiting human approval",
        "Label / Manufacturer / Country — reference entities (confirmed by reviewers)",
        "IngestionJob — background connector job tracking",
        "AuditLog — append-only immutable action trail",
        "User — application user with role-based access",
    ]
    for item in er_desc:
        _bullet(doc, item)

    _h2(doc, "3.2 Core Table Reference")

    tables_data = [
        # (table_name, description, key_columns)
        ("users", "Application users — reviewers, admins, service accounts",
         "id · username · email · password_hash (Argon2id) · role · is_active · "
         "mfa_secret · failed_login_count · locked_until"),

        ("tokens", "Core token object — one row per unique physical release",
         "id · barcode · catalog_number · matrix_runout (plain text, side A) · "
         "matrix_runout_b · matrix_runout_parts (JSONB) · matrix_runout_b_parts (JSONB) · "
         "media_format · status · title · artist · year · country_id · label_id · "
         "manufacturer_id · discogs_release_id · musicbrainz_release_id · "
         "current_revision_id · extra (JSONB)"),

        ("token_revisions", "Append-only revision log; stores complete attribute snapshot",
         "id · token_id · revision_number · source · data (JSONB) · diff (JSONB) · "
         "confidence · created_by · ingestion_job_id"),

        ("media_objects", "Images stored in MinIO; keyed by SHA-256 content hash",
         "id · token_id · content_hash · bucket · object_key · mime_type · "
         "size_bytes · width_px · height_px · region · source_url · metadata (JSONB)"),

        ("ocr_artifacts", "OCR text per image region; includes symbol candidates",
         "id · media_object_id · engine · region · raw_text · detected_language · "
         "translated_text · confidence · symbol_candidates (JSONB) · metadata (JSONB)"),

        ("symbols", "Canonical graphical runout / dead-wax symbol registry",
         "id · slug (immutable) · name · category · description · unicode_approx · "
         "taxonomy_level (1–5) · region_scope · is_confirmed · metadata (JSONB)"),

        ("symbol_variants", "Visual variants of a canonical symbol",
         "id · symbol_id · variant_key · description · reference_image_key"),

        ("token_symbols", "FK index: symbol↔token position for fast joins",
         "id · token_id · symbol_id · position · side (a|b)"),

        ("labels", "Record label reference; confirmed by human reviewer",
         "id · name · name_normalised · country_id · discogs_id · musicbrainz_id · "
         "is_confirmed · metadata (JSONB)"),

        ("manufacturers", "Pressing plant / manufacturer reference",
         "id · name · name_normalised · country_id · plant_code · is_confirmed"),

        ("countries", "ISO 3166-1 country seed data",
         "id · alpha2 · alpha3 · name · numeric_code"),

        ("ingestion_jobs", "Background connector job tracking",
         "id · connector_name · status · payload (JSONB) · result (JSONB) · "
         "error_message · attempt_count · started_at · completed_at"),

        ("review_items", "Human-review queue; all AI proposals land here first",
         "id · token_id · revision_id · status · reason · priority · "
         "details (JSONB) · assigned_to · resolved_at · resolution_comment"),

        ("audit_log", "Immutable append-only action log (BIGINT PK, no updates)",
         "id · timestamp · user_id · action · entity_type · entity_id · "
         "detail (JSONB) · ip_address · request_id"),
    ]

    for tbl_name, tbl_desc, tbl_cols in tables_data:
        _h3(doc, tbl_name)
        _body(doc, tbl_desc)
        p = doc.add_paragraph()
        _add_run(p, "Columns: ", bold=True, colour=C_NAVY, size=9)
        _add_run(p, tbl_cols, mono=True, size=9, colour=C_SLATE)
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(6)

    _h2(doc, "3.3 Enumerations")
    _simple_table(doc,
        ["Enum Name", "Values"],
        [
            ["media_format",         "vinyl · cd"],
            ["token_status",         "draft · active · merged · archived"],
            ["revision_source",      "ingestion · vision · ocr · human · llm · import"],
            ["review_status",        "pending · in_progress · approved · rejected · deferred"],
            ["review_reason",        "low_confidence · conflict · novel_entity · anomaly · manual"],
            ["ingestion_job_status", "queued · running · completed · failed · cancelled"],
            ["ocr_engine",           "tesseract · azure · aws_textract · manual"],
            ["image_region",         "label_a · label_b · obi_front/back/spine · runout_a/b · matrix · cover_front/back · sleeve_inner · disc_surface · other"],
            ["user_role",            "admin · reviewer · viewer · service"],
            ["symbol_category",      "pressing_plant_mark · engineer_mark · label_logo · cut_type · certification · other"],
        ],
        col_widths=[4.5, 10.5],
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # 4. VISION & OCR PIPELINE
    # ═══════════════════════════════════════════════════════════════════════
    _h1(doc, "4 · Vision & OCR Pipeline")

    _h2(doc, "4.1 Pipeline Overview")
    steps_vision = [
        ("1", "Image arrives",        "Uploaded by user or fetched from Discogs / CoverArtArchive"),
        ("2", "MinIO store",           "SHA-256 hash computed → deduplicated upload to MinIO bucket"),
        ("3", "OCR extraction",        "Tesseract (or cloud: Azure / AWS Textract) per image region"),
        ("4", "Translation",           "OCR text translated to British English via local LLM"),
        ("5", "Vision transcription",  "HybridVision calls Ollama (LLaVA / Qwen2-VL) with task prompt"),
        ("6", "JSON parsing",          "Structured response validated against expected schema"),
        ("7", "Candidate matching",    "Trigram search + exact match against token table"),
        ("8", "Review queue",          "Results written to review_items — never auto-applied"),
    ]
    _simple_table(doc,
        ["Step", "Stage", "Detail"],
        steps_vision,
        col_widths=[1.2, 4.0, 9.8],
    )

    _h2(doc, "4.2 Vision Backends")
    _body(doc,
        "The HybridVision adapter tries the primary backend first; on timeout or error "
        "it falls through to the fallback.  Both backends are transparent to the caller.")
    for backend, detail in [
        ("OllamaVisionBackend (primary)",
         "Sends base64-encoded image + prompt to http://ollama:11434/api/chat.  "
         "Models: LLaVA 1.6, Qwen2-VL.  Timeout: 120 s.  Runs entirely on-premises."),
        ("AnthropicVisionBackend (fallback)",
         "Uses the Anthropic Messages API (claude-3-5-sonnet or configured model).  "
         "Requires ANTHROPIC_API_KEY.  Invoked only when Ollama fails or confidence is too low."),
    ]:
        p = doc.add_paragraph()
        _add_run(p, backend + " — ", bold=True, colour=C_NAVY, size=10)
        _add_run(p, detail, colour=C_SLATE, size=10)
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(4)

    _h2(doc, "4.3 Prompt Templates")
    _simple_table(doc,
        ["Function", "Region", "Key Output Fields"],
        [
            ["label_prompt()",               "label_a / label_b",     "label_name, catalog_number, artist, title, side, speed_rpm, country, year"],
            ["obi_prompt()",                 "obi_front / back / spine", "japanese_title, romanised_title, english_title, catalog_number, price, obi_type"],
            ["runout_prompt()",              "runout_a / runout_b / matrix", "matrix_number, stamper_code, sid_codes, lacquer_cut_info, pressing_plant_hint, symbol_detections"],
            ["symbol_identification_prompt()","runout (re-run)",       "symbols[] with slug_suggestion, unicode_approx, description, application, confidence"],
        ],
        col_widths=[5.0, 4.0, 6.0],
    )

    _h2(doc, "4.4 Symbol Detections in Runout Output")
    _body(doc,
        "When the vision model encounters a non-alphanumeric graphical mark in the "
        "dead-wax area, it emits a symbol_detections entry rather than embedding the "
        "mark as plain text.  Each entry has the following shape:")
    _code(doc, '{"slug_suggestion": "emi-triangle",  // null if unknown')
    _code(doc, ' "unicode_approx":  "△",')
    _code(doc, ' "description":     "Upward triangle stamped into wax, ~3 mm",')
    _code(doc, ' "application":     "stamped",')
    _code(doc, ' "confidence":      0.92}')

    _callout(doc, "Design invariant:",
        "Vision proposals are advisory only.  All detections are written to the "
        "review queue; they are never applied to the token or symbol tables without "
        "explicit human approval.", colour=C_TEAL)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # 5. SYMBOL REGISTRY
    # ═══════════════════════════════════════════════════════════════════════
    _h1(doc, "5 · Symbol Registry")

    _h2(doc, "5.1 Purpose")
    _body(doc,
        "Runout / dead-wax inscriptions on vinyl records mix plain text with graphical "
        "symbols — pressed triangles, stamped stars, etched circles — that identify "
        "pressing plants, mastering engineers, label certifications, and cutting systems.  "
        "The symbol registry gives each mark a stable, human-readable slug and enables "
        "indexed queries without scanning JSONB arrays.")

    _h2(doc, "5.2 Taxonomy — Five Levels of Rarity")
    _simple_table(doc,
        ["Level", "Frequency", "Description", "Examples"],
        [
            ["1", "Very common",
             "Core text content handled as plain text",
             "Matrix numbers (XZAL-9067), side codes (A/B), stamper codes (1A), basic plant text (EMI, CBS)"],
            ["2", "Common graphical symbols — seeded at install",
             "Appears on thousands of records",
             "EMI △ (UK), PRS ▽ (UK), Capitol ☆ (US), Decca ◆/◈/✤ (US), Porky/Pecko marks, Pye Studios △M"],
            ["3", "Regional / label-specific",
             "Familiar to specialists; import from reference data",
             "Columbia plant codes (Ƨ/T/P/G), Allied (a/Q), Sterling Sound stamp, Masterdisk stamp, Wakefield tulip"],
            ["4", "Specialist / Vintage",
             "Rare but identifiable to experts",
             "Western Electric □/◇, Lindström £/ℒ, Japanese JIS 〄, Nigerian/Jamaican plant marks"],
            ["5", "Edge cases",
             "One in a thousand — entered manually",
             "Direct-cut indicators, unique handwritten engineer marks, test-pressing one-offs"],
        ],
        col_widths=[1.2, 3.0, 4.5, 6.3],
    )

    _h2(doc, "5.3 Seeded Symbols (Levels 2–4)")
    _simple_table(doc,
        ["Slug", "Name", "Category", "Region", "Level"],
        [
            ["emi-triangle",              "EMI Pressing Triangle",            "pressing_plant_mark", "UK",      "2"],
            ["prs-triangle-down",         "PRS Downward Triangle ▽",          "certification",       "UK",      "2"],
            ["pye-triangle",              "Pye Studios Engineer Triangle",    "engineer_mark",       "UK",      "2"],
            ["porky-prime-cut",           "Porky Prime Cut",                  "engineer_mark",       "UK",      "2"],
            ["pecko-duck",                "Pecko Duck (alt. Peckham mark)",   "engineer_mark",       "UK",      "2"],
            ["decca-circle",              "Decca / London Circle",            "pressing_plant_mark", "UK",      "2"],
            ["sonic-arts-logo",           "Sonic Arts Logo ▭◯▭",             "label_logo",          "UK",      "2"],
            ["capitol-la-star",           "Capitol Los Angeles Star ☆",       "pressing_plant_mark", "US",      "2"],
            ["decca-us-gloversville",     "MCA/Decca Gloversville ✤",         "pressing_plant_mark", "US",      "2"],
            ["decca-us-pinckneyville",    "MCA/Decca Pinckneyville ◆",        "pressing_plant_mark", "US",      "2"],
            ["decca-us-richmond",         "MCA/Decca Richmond ◈",             "pressing_plant_mark", "US",      "2"],
            ["sheffield-lab-delta",       "Sheffield Lab △####",              "pressing_plant_mark", "US",      "2"],
            ["sterling-sound",            "Sterling Sound stamp",             "engineer_mark",       "US",      "3"],
            ["masterdisk",                "Masterdisk stamp",                 "engineer_mark",       "US",      "3"],
            ["columbia-santa-maria",      "Columbia Santa Maria (Ƨ)",         "pressing_plant_mark", "US",      "3"],
            ["columbia-terre-haute",      "Columbia Terre Haute (T/CT/CTH)", "pressing_plant_mark", "US",      "3"],
            ["columbia-pitman",           "Columbia Pitman (P)",              "pressing_plant_mark", "US",      "3"],
            ["columbia-carrollton",       "Columbia Carrollton (G/G1)",       "pressing_plant_mark", "US",      "3"],
            ["capitol-jacksonville",      "Capitol Jacksonville (0/())",      "pressing_plant_mark", "US",      "3"],
            ["capitol-winchester",        "Capitol Winchester (—◁)",          "pressing_plant_mark", "US",      "3"],
            ["capitol-scranton-iam",      "Capitol Scranton (IAM △)",         "pressing_plant_mark", "US",      "3"],
            ["allied-record-a",           "Allied Record (a/Q)",              "pressing_plant_mark", "US",      "3"],
            ["wakefield-tulip",           "Wakefield Manufacturing tulip",    "pressing_plant_mark", "US",      "3"],
            ["western-electric-blumlein-square","WE Blumlein Square □",      "cut_type",            "—",       "4"],
            ["western-electric-diamond",  "WE Diamond ◇ (1C/1D)",            "cut_type",            "—",       "4"],
            ["lindstrom-pound",           "Lindström System £/ℒ",             "cut_type",            "Europe",  "4"],
            ["japanese-jis",              "Japanese JIS Mark 〄",              "certification",       "Japan",   "4"],
        ],
        col_widths=[4.5, 4.5, 3.5, 1.8, 1.2],
    )

    _h2(doc, "5.4 Parts Array Format")
    _body(doc,
        "Once symbols are confirmed, the plain-text matrix_runout field is supplemented "
        "with a structured parts array stored in matrix_runout_parts (JSONB).  "
        "Each element is either a text fragment or a resolved symbol reference:")
    _code(doc, '[ {"t": "text", "v": "A1 "},')
    _code(doc, '  {"t": "sym",  "slug": "emi-triangle", "id": "<uuid>"},')
    _code(doc, '  {"t": "text", "v": " XZA 1234-1"} ]')

    _h2(doc, "5.5 Symbol Helpers")
    for fn, sig, desc in [
        ("render_parts_to_text",
         'render_parts_to_text(parts, *, symbols: dict[str,str]) → str',
         'Converts a parts array back to a plain-text string using a slug→display mapping.  '
         'Unknown slugs render as [slug] so nothing is silently dropped.'),
        ("extract_symbol_ids",
         'extract_symbol_ids(parts) → list[tuple[str, int]]',
         'Returns (uuid, position) pairs for every symbol entry, used to rebuild the '
         'token_symbols FK index after a parts array update.'),
    ]:
        _h3(doc, fn)
        _code(doc, sig)
        _body(doc, desc)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # 6. INGESTION CONNECTORS
    # ═══════════════════════════════════════════════════════════════════════
    _h1(doc, "6 · Ingestion Connectors")

    _h2(doc, "6.1 Architecture")
    _body(doc,
        "Connectors run inside the dedicated worker container.  Each connector extends "
        "BaseConnector and is governed by three independent resilience mechanisms:")
    for mech, detail in [
        ("Token-bucket rate limiter",
         "Limits outbound request rate per connector (configurable; Discogs default 1 req/s)."),
        ("Circuit breaker",
         "Opens after N consecutive failures; automatically recovers after a timeout.  "
         "Prevents hammering a failing upstream during outages."),
        ("Exponential backoff retry",
         "Re-attempts failed fetches up to max_attempts (default 3–5) with configurable backoff factor."),
    ]:
        p = doc.add_paragraph()
        _add_run(p, mech + " — ", bold=True, colour=C_NAVY, size=10)
        _add_run(p, detail, colour=C_SLATE, size=10)
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(4)

    _h2(doc, "6.2 Available Connectors")
    _simple_table(doc,
        ["Connector", "Source", "Operations", "Rate Limit"],
        [
            ["Discogs",      "discogs.com API",         "fetch_release(id), search_releases(query)", "1 req/s (3600/hr)"],
            ["MusicBrainz",  "musicbrainz.org API",     "fetch_release(id), search_releases(query)", "1 req/s; User-Agent required"],
        ],
        col_widths=[3.0, 4.0, 5.5, 2.5],
    )

    _h2(doc, "6.3 Redis Job Queue")
    _body(doc,
        "Jobs are stored as JSON in Redis lists.  The BLMOVE command atomically moves "
        "a job from pending to processing, providing exactly-once delivery even if the "
        "worker crashes mid-job.")
    _simple_table(doc,
        ["Redis Key", "Purpose"],
        [
            ["mediacat:jobs:pending",          "FIFO queue of jobs awaiting processing"],
            ["mediacat:jobs:processing",       "In-flight jobs (dequeued but not yet complete)"],
            ["mediacat:jobs:dead",             "Jobs that exhausted all retry attempts"],
            ["mediacat:jobs:processing_times", "Hash of job_id → dequeue timestamp (staleness detection)"],
        ],
        col_widths=[6.0, 9.0],
    )

    _h2(doc, "6.4 Job Lifecycle")
    lifecycle = [
        "Scheduler enqueues Job(connector, action, payload)",
        "Worker calls BLMOVE pending → processing",
        "Connector fetches data + image URLs",
        "Storage pipeline downloads + deduplicates images in MinIO",
        "OCR extracts text per image region",
        "Vision model transcribes label / OBI / runout",
        "Rule engine decodes matrix codes (OPA or Python fallback)",
        "Token created or revision appended",
        "Low-confidence or conflicting results → ReviewItem queued",
        "LREM removes job from processing (success)",
        "On crash: stale-job reaper re-enqueues after 600 s",
    ]
    for i, step in enumerate(lifecycle):
        _bullet(doc, f"{i+1}.  {step}")

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # 7. REVIEW QUEUE
    # ═══════════════════════════════════════════════════════════════════════
    _h1(doc, "7 · Review Queue")

    _h2(doc, "7.1 Triggers")
    _simple_table(doc,
        ["Trigger", "Reason Code", "Description"],
        [
            ["Vision / OCR confidence below threshold", "low_confidence", "Vision model returns confidence < 0.7"],
            ["Multiple sources disagree",               "conflict",       "Discogs and MusicBrainz return different values for the same field"],
            ["Unknown label or manufacturer",           "novel_entity",   "Pipeline proposes an entity not in the reference tables"],
            ["LLM anomaly detection",                   "anomaly",        "Drift detector flags unexpected field values or schema changes"],
            ["Manual raise",                            "manual",         "Reviewer triggers re-review of an existing token"],
        ],
        col_widths=[5.0, 3.5, 6.5],
    )

    _h2(doc, "7.2 Review Item States")
    for state, arrow, desc in [
        ("pending",     "→",  "Newly queued; waiting for a reviewer to pick up"),
        ("in_progress", "→",  "A reviewer has opened the item"),
        ("approved",    "",   "Revision accepted; token updated; audit logged"),
        ("rejected",    "",   "Revision discarded; reason captured in resolution_comment"),
        ("deferred",    "",   "Held for later; may be reassigned or escalated"),
    ]:
        p = doc.add_paragraph()
        _add_run(p, state, mono=True, bold=True, size=10, colour=C_TEAL)
        _add_run(p, f"  —  {desc}", size=10, colour=C_SLATE)
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(3)

    _h2(doc, "7.3 Workflow")
    for step, title, detail in [
        ("1.", "Browse the queue",  "Navigate to /reviews.  Items are sorted by priority (desc) then age."),
        ("2.", "Open a review",     "Click any item to see the token's current values, the proposed revision, and the diff."),
        ("3.", "Examine evidence",  "View source images, OCR text, and vision model confidence scores."),
        ("4.", "Approve or Reject", "POST to /reviews/{id}/approve or /reviews/{id}/reject with an optional comment.  "
                                    "Both actions are CSRF-protected and logged to audit_log."),
        ("5.", "Token updated",     "On approval, the revision is applied to the token's denormalised fields and "
                                    "current_revision_id is advanced."),
    ]:
        p = doc.add_paragraph()
        _add_run(p, step + " " + title + " — ", bold=True, colour=C_NAVY, size=10)
        _add_run(p, detail, colour=C_SLATE, size=10)
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(4)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # 8. AUTHENTICATION & SECURITY
    # ═══════════════════════════════════════════════════════════════════════
    _h1(doc, "8 · Authentication & Security")

    _h2(doc, "8.1 Authentication Flow")
    for step, detail in [
        ("1. Password hashing",   "Argon2id (time_cost=3, memory=64 MiB, parallelism=2, hash_len=32 B, salt=16 B)."),
        ("2. Login rate limiting","In-memory lockout; max 10 failures per username or IP in 15-minute window."),
        ("3. Account lockout",    "User.locked_until prevents further attempts after threshold reached."),
        ("4. Session creation",   "Signed cookie (itsdangerous.TimestampSigner); payload = user_id|role|nonce; 24 h TTL."),
        ("5. CSRF protection",    "Per-session HMAC-SHA256 token validated from X-CSRF-Token header on all mutating requests."),
        ("6. MFA (scaffold)",     "User.mfa_secret stores TOTP secret (encrypted at rest); not yet wired into login flow."),
    ]:
        p = doc.add_paragraph()
        _add_run(p, step + " — ", bold=True, colour=C_NAVY, size=10)
        _add_run(p, detail, colour=C_SLATE, size=10)
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(4)

    _h2(doc, "8.2 User Roles")
    _simple_table(doc,
        ["Role", "Capabilities"],
        [
            ["admin",    "Full access: user management, review, browse, all admin functions"],
            ["reviewer", "Approve / reject review items; browse tokens"],
            ["viewer",   "Read-only: browse tokens, view revisions"],
            ["service",  "Machine-to-machine: ingestion jobs, OCR, vision pipeline"],
        ],
        col_widths=[3.0, 12.0],
    )

    _h2(doc, "8.3 Database Roles (Least Privilege)")
    _simple_table(doc,
        ["DB Role", "Permissions"],
        [
            ["postgres",           "Superuser — used for initial setup only"],
            ["mediacat_migrator",  "Owns schema; runs Alembic; CREATE/DROP/ALTER"],
            ["mediacat_app",       "INSERT/SELECT/UPDATE on app tables; no DDL; no UPDATE/DELETE on audit_log"],
            ["mediacat_readonly",  "SELECT only — for reporting and analytics"],
        ],
        col_widths=[4.0, 11.0],
    )

    _h2(doc, "8.4 Security Headers")
    for hdr, val in [
        ("Content-Security-Policy", "Strict; no inline scripts; form-action restricted"),
        ("X-Content-Type-Options",  "nosniff"),
        ("X-Frame-Options",         "DENY"),
        ("Referrer-Policy",         "strict-origin-when-cross-origin"),
        ("Permissions-Policy",      "camera=(), microphone=(), geolocation=()"),
        ("Strict-Transport-Security","max-age=63072000 (2 years)"),
        ("Server / X-Powered-By",   "Removed — no server fingerprinting"),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(2)
        _add_run(p, hdr + ": ", mono=True, bold=True, size=9, colour=C_NAVY)
        _add_run(p, val, size=9, colour=C_SLATE)

    _h2(doc, "8.5 Network Security")
    for item in [
        "Caddy reverse proxy terminates TLS; app only binds to 127.0.0.1:8000.",
        "Frontend (public) and backend (internal) Docker networks are isolated.",
        "All secrets injected via Docker secret files (/run/secrets/); never in environment variables.",
        "All parameterised queries via SQLAlchemy — no raw SQL, no injection risk.",
        "Input validated by Pydantic models at every external boundary.",
        "Jinja2 templates auto-escape HTML — no XSS via template output.",
    ]:
        _bullet(doc, item)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # 9. OBJECT STORAGE
    # ═══════════════════════════════════════════════════════════════════════
    _h1(doc, "9 · Object Storage (MinIO)")

    _h2(doc, "9.1 Buckets")
    _simple_table(doc,
        ["Bucket", "Contents", "Key Pattern"],
        [
            ["media-originals", "Raw images fetched from Discogs, CoverArtArchive, user upload", "<sha256>.<ext>"],
            ["ocr-artifacts",   "OCR extraction results, parsing artefacts",                    "<sha256>.json"],
            ["backups",         "Daily pg_dump snapshots and database archives",                 "YYYY-MM-DD/<dump>.sql.gz"],
        ],
        col_widths=[4.0, 7.0, 4.0],
    )

    _h2(doc, "9.2 Content-Hash Deduplication")
    _body(doc,
        "Every image is hashed with SHA-256 before upload.  If a file with the same "
        "hash already exists in MinIO, no upload occurs — the existing object is referenced.  "
        "This means the same album cover fetched from Discogs and uploaded manually by "
        "a user occupies exactly one object in storage.  The MediaObject row always points "
        "to the existing hash regardless of source.")

    _h2(doc, "9.3 Image Validation")
    for check in [
        "MIME type whitelist: JPEG, PNG, TIFF, WebP, GIF, BMP only.",
        "Pillow opens the image to verify it is structurally valid.",
        "Maximum pixel count enforced (178.9 MP) — prevents decompression bomb attacks.",
        "Width and height extracted and stored in media_objects for layout hints.",
    ]:
        _bullet(doc, check)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # 10. CONFIGURATION
    # ═══════════════════════════════════════════════════════════════════════
    _h1(doc, "10 · Configuration Reference")

    _h2(doc, "10.1 app.yaml — Key Sections")
    sections_cfg = [
        ("app",          "name, environment (dev|staging|prod), log_level"),
        ("server",       "host, port (default 8000)"),
        ("security",     "session_secret (resolved from secret file), cookie_secure, login lockout thresholds"),
        ("postgres",     "host, port, user, database, password_file"),
        ("object_store", "endpoint, access_key, secret_key_file, default_bucket"),
        ("redis",        "url template, password_file"),
        ("rule_engine",  "backend (opa|local), opa_url"),
        ("vision",       "backend (hybrid), primary (ollama), ollama_url, fallback (anthropic)"),
        ("llm",          "backend (hybrid), primary (ollama), fallback (anthropic)"),
        ("feature_flags","vision_local, llm_local, api_fallback — toggle per environment"),
    ]
    _simple_table(doc,
        ["Section", "Key Settings"],
        sections_cfg,
        col_widths=[3.5, 11.5],
    )

    _h2(doc, "10.2 .env — Key Variables")
    _simple_table(doc,
        ["Variable", "Description", "Example"],
        [
            ["MEDIACAT_ENV",          "Runtime environment",                           "dev"],
            ["MEDIACAT_DATA_ROOT",    "Host path for volumes and secrets",             "~/data/mediacat"],
            ["HTTP_BIND",             "Port binding for Caddy HTTP",                   "0.0.0.0:80"],
            ["HTTPS_BIND",            "Port binding for Caddy HTTPS",                  "0.0.0.0:443"],
            ["PUBLIC_HOSTNAME",       "External hostname (for Let's Encrypt)",         "mediacat.example.com"],
            ["MEDIACAT_DEV_ADMIN_PASSWORD", "Seed password for dev admin user",        "(choose a strong password)"],
            ["TZ",                    "Timezone for containers",                        "Europe/London"],
        ],
        col_widths=[5.0, 5.5, 4.5],
    )

    _h2(doc, "10.3 Docker Secrets")
    _body(doc,
        "All secret values are injected via files in ${MEDIACAT_DATA_ROOT}/secrets/.  "
        "Each file contains only the secret value with no trailing newline.")
    _simple_table(doc,
        ["Secret File", "Used By"],
        [
            ["postgres_app_password",   "PostgreSQL connection string for the app role"],
            ["minio_root_password",     "MinIO admin credentials"],
            ["redis_password",          "Redis AUTH password"],
            ["session_secret",          "HMAC key for session cookie signing"],
            ["discogs_token",           "Discogs API personal access token (optional)"],
        ],
        col_widths=[5.5, 9.5],
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # 11. API ENDPOINTS
    # ═══════════════════════════════════════════════════════════════════════
    _h1(doc, "11 · API Endpoints")

    _body(doc,
        "MediaCat uses FastAPI with Jinja2-rendered HTML responses.  The primary "
        "interface is the web UI; HTMX is used for partial updates.  All mutating "
        "endpoints require a valid session and an X-CSRF-Token header.")

    _h2(doc, "11.1 Public (No Auth Required)")
    _simple_table(doc,
        ["Method", "Path", "Description"],
        [
            ["GET",  "/healthz",  "Liveness probe — returns {\"status\":\"ok\"}"],
            ["GET",  "/readyz",   "Readiness probe — returns {\"status\":\"ok\"}"],
            ["GET",  "/login",    "Render login form with CSRF token"],
            ["POST", "/login",    "Authenticate — creates session cookie, redirects to /"],
            ["GET",  "/logout",   "Clear session cookie, redirect to /login"],
        ],
        col_widths=[1.5, 3.5, 10.0],
    )

    _h2(doc, "11.2 Dashboard & Browse (Auth Required)")
    _simple_table(doc,
        ["Method", "Path", "Description"],
        [
            ["GET", "/",                          "Main dashboard — pending count, stats"],
            ["GET", "/tokens",                    "Browse tokens; query params: q, media, page"],
            ["GET", "/tokens/{token_id}",         "Single token detail with revision history"],
        ],
        col_widths=[1.5, 5.5, 8.0],
    )

    _h2(doc, "11.3 Review Queue (Reviewer+ Role)")
    _simple_table(doc,
        ["Method", "Path", "Description"],
        [
            ["GET",  "/reviews",                       "List review items; params: status, page"],
            ["GET",  "/reviews/{review_id}",           "Single review detail with diff view"],
            ["POST", "/reviews/{review_id}/approve",   "Approve revision; form: comment"],
            ["POST", "/reviews/{review_id}/reject",    "Reject revision; form: comment"],
        ],
        col_widths=[1.5, 5.5, 8.0],
    )

    _h2(doc, "11.4 User Management (Admin Role)")
    _simple_table(doc,
        ["Method", "Path", "Description"],
        [
            ["GET",  "/users",      "List all users"],
            ["GET",  "/users/new",  "Render create-user form"],
            ["POST", "/users/new",  "Create user; form: username, email, password, role"],
        ],
        col_widths=[1.5, 4.0, 9.5],
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # 12. MIGRATIONS
    # ═══════════════════════════════════════════════════════════════════════
    _h1(doc, "12 · Alembic Migrations")

    _body(doc,
        "Migrations are sequential and deterministic.  All constraint names follow a "
        "naming convention (pk_, fk_, uq_, ix_) to avoid name collisions.  "
        "Run migrations as the mediacat_migrator role.")
    _code(doc, "# Apply all pending migrations")
    _code(doc, "docker compose exec app alembic upgrade head")
    _code(doc, "")
    _code(doc, "# Roll back one migration")
    _code(doc, "docker compose exec app alembic downgrade -1")
    _code(doc, "")
    _code(doc, "# Show current revision")
    _code(doc, "docker compose exec app alembic current")

    _h2(doc, "Migration 0001 — initial_schema (2026-04-17)")
    _body(doc, "Creates the complete initial schema:")
    for item in [
        "9 PostgreSQL ENUM types",
        "Extensions: uuid-ossp, pg_trgm, btree_gist",
        "11 application tables: users, countries, labels, manufacturers, tokens, "
        "token_revisions, media_objects, ocr_artifacts, ingestion_jobs, review_items, audit_log",
        "Trigram GIN indexes on labels.name_normalised and manufacturers.name_normalised",
        "Trigram GIN indexes on tokens.title and tokens.artist",
        "Partial index on review_items WHERE status = 'pending'",
        "REVOKE UPDATE, DELETE ON audit_log FROM mediacat_app",
    ]:
        _bullet(doc, item)

    _h2(doc, "Migration 0002 — symbols (2026-04-22)")
    _body(doc, "Adds the symbol registry and extends runout columns:")
    for item in [
        "New ENUM: symbol_category (6 values)",
        "New table: symbols (slug, name, category, unicode_approx, taxonomy_level, region_scope)",
        "New table: symbol_variants (variant_key, reference_image_key per symbol)",
        "New table: token_symbols (FK index: token ↔ symbol position)",
        "New columns on tokens: matrix_runout_b (TEXT), matrix_runout_parts (JSONB), matrix_runout_b_parts (JSONB)",
        "New column on ocr_artifacts: symbol_candidates (JSONB)",
        "Seed data: 26 confirmed symbols at taxonomy levels 2–4 (UK, US, European, Japanese marks)",
    ]:
        _bullet(doc, item)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # 13. MAKE TARGETS
    # ═══════════════════════════════════════════════════════════════════════
    _h1(doc, "13 · Make Targets")

    _simple_table(doc,
        ["Target", "Category", "Description"],
        [
            ["bootstrap",    "Setup",      "Run Ubuntu bootstrap script (sudo required); installs build tools, Tesseract"],
            ["data-init",    "Setup",      "Create/verify persistent data directory tree on host"],
            ["venv",         "Setup",      "Create local Python virtualenv at .venv/"],
            ["setup",        "Setup",      "venv + install deps (uv sync) + install pre-commit hooks"],
            ["deps-sync",    "Setup",      "Re-sync Python dependencies from pyproject.toml"],
            ["lint",         "Quality",    "Ruff linter"],
            ["format",       "Quality",    "Ruff formatter + import sorter"],
            ["typecheck",    "Quality",    "Mypy strict static type checking"],
            ["security",     "Quality",    "Bandit (code) + pip-audit (dependencies)"],
            ["test",         "Quality",    "Full pytest suite with coverage"],
            ["test-fast",    "Quality",    "Pytest excluding slow / integration markers"],
            ["docs-api",     "Docs",       "Generate API reference with pdoc → docs/reference/"],
            ["docs-build",   "Docs",       "Build MkDocs site to site/"],
            ["docs-serve",   "Docs",       "Live-preview docs on http://127.0.0.1:8800"],
            ["up",           "Docker",     "Start all 8 Docker services (background)"],
            ["down",         "Docker",     "Stop services; data volumes preserved"],
            ["restart",      "Docker",     "Restart the stack"],
            ["logs",         "Docker",     "Tail logs from all services"],
            ["ps",           "Docker",     "docker compose ps — show service status"],
            ["config-check", "Docker",     "Validate docker-compose configuration"],
            ["clean",        "Cleanup",    "Remove build artefacts and caches"],
            ["distclean",    "Cleanup",    "clean + remove .venv/"],
        ],
        col_widths=[3.0, 2.0, 10.0],
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # 14. OPEN ITEMS
    # ═══════════════════════════════════════════════════════════════════════
    _h1(doc, "14 · Open Items & Roadmap")

    _h2(doc, "14.1 Security Defects (Tracked)")
    defects = [
        ("DEF-001", "MEDIUM", "Wire sanitise() into translation pipeline",
         "OCR text passed to LLM for translation without sanitisation — "
         "malicious text in a scanned label could inject instructions into the LLM prompt.  "
         "Fix: apply sanitise() before truncation, before LLM call.",
         "1 hour"),
        ("DEF-002", "LOW", "Migrate login rate limiter to Redis",
         "LoginRateLimiter is an in-memory dict — state is lost on app restart and "
         "not shared across multiple app instances.  Redis backend provides persistence "
         "and cluster safety.",
         "0.5 day"),
        ("DEF-003", "MEDIUM", "REVOKE UPDATE/DELETE on audit_log at DB level",
         "Currently enforced via OPA policy only.  A direct DB connection (e.g. a "
         "compromised session) could still delete audit rows.  "
         "Fix: execute REVOKE UPDATE, DELETE ON audit_log FROM mediacat_app; after migration.",
         "1 hour"),
        ("DEF-004", "LOW", "detect_is_english false positives",
         "Simple substring-match heuristic may misclassify non-English text.  "
         "Replace with a proper language detection library.",
         "0.5 day"),
        ("DEF-005", "LOW", "BaseHTTPMiddleware body-streaming overhead",
         "Starlette BaseHTTPMiddleware buffers the full request/response body.  "
         "Migrate to pure ASGI middleware for better performance on large image uploads.",
         "1–2 days"),
    ]
    for ref, sev, title, detail, effort in defects:
        sev_colour = C_RED if sev == "MEDIUM" else C_AMBER
        _callout(doc, f"{ref} [{sev}] — {title}", f"{detail}  (Effort: {effort})", colour=sev_colour)

    _h2(doc, "14.2 Architecture Decisions (ADRs)")
    _simple_table(doc,
        ["ADR", "Decision", "Status"],
        [
            ["ADR-0001", "Media scope: vinyl + CD only (v1)",                             "Accepted"],
            ["ADR-0002", "Hybrid AI: local Ollama primary, API fallback",                 "Accepted"],
            ["ADR-0003", "Rule engine: OPA primary, Python local fallback",               "Accepted"],
            ["ADR-0004", "Data/code separation via host bind mounts",                     "Accepted"],
            ["ADR-0005", "Web stack: FastAPI + Jinja2 + HTMX (no JS framework)",         "Accepted"],
        ],
        col_widths=[2.0, 10.5, 2.5],
    )

    _h2(doc, "14.3 Near-term Roadmap")
    for pri, item in [
        ("P1", "Wire DEF-001 sanitise() into translation pipeline"),
        ("P1", "Enforce DEF-003 audit_log REVOKE at DB level"),
        ("P2", "DEF-002 Redis-backed login rate limiter"),
        ("P2", "Phase 2 symbol pipeline: focused re-run with slug hints"),
        ("P2", "LLM comparison / anomaly detection module (llm/tasks.py)"),
        ("P3", "MFA TOTP login flow (scaffold already in User model)"),
        ("P3", "CLIP embedding similarity search for symbol Phase 3"),
        ("P3", "Horizontal scaling: multiple worker containers on same Redis queue"),
    ]:
        p = doc.add_paragraph()
        _add_run(p, f"[{pri}]  ", bold=True,
                 colour=C_RED if pri == "P1" else (C_AMBER if pri == "P2" else C_TEAL),
                 size=10)
        _add_run(p, item, colour=C_SLATE, size=10)
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(3)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # 15. USAGE GUIDE
    # ═══════════════════════════════════════════════════════════════════════
    _h1(doc, "15 · Usage Guide")

    _h2(doc, "15.1 Daily Workflow — Cataloguing a New Record")
    workflow = [
        ("1. Add the token",
         'Navigate to /tokens and click "New Token".  Enter the basic details: '
         'media format (vinyl/CD), barcode or catalogue number if known.  '
         'The token is created in Draft status.'),
        ("2. Upload images",
         'On the token detail page, upload scans for each region: label side A, '
         'label side B, runout A, runout B, cover front/back.  Use high-resolution '
         'scans (≥ 600 DPI) for runout areas to help the vision model detect symbols.'),
        ("3. Trigger ingestion",
         'Click "Fetch from Discogs" or "Fetch from MusicBrainz" to pull metadata.  '
         'The system queues a job; the worker fetches, OCRs, and runs vision in the background.'),
        ("4. Review proposals",
         'Navigate to /reviews.  Each AI proposal appears as a pending review item.  '
         'Check the diff (proposed vs current), inspect confidence scores, and '
         'approve or reject each field update.'),
        ("5. Confirm symbols",
         'For runout images, symbol_detections appear in the review item details.  '
         'If a slug suggestion is shown (e.g. emi-triangle), verify it visually and '
         'approve.  If no slug was suggested, look up the symbol in the registry and '
         'assign the correct slug before approving.'),
        ("6. Activate the token",
         'Once you are satisfied with the data, change the token status from Draft '
         'to Active on the token detail page.'),
    ]
    for step, detail in workflow:
        _h3(doc, step)
        _body(doc, detail)

    _h2(doc, "15.2 Searching for Tokens")
    _body(doc,
        "The token browser at /tokens supports full-text search across title and artist "
        "(backed by PostgreSQL trigram similarity) and filtering by media format.  "
        "The search is tolerant of minor spelling errors — a query for 'Beatls' will "
        "still surface Beatles releases.")
    for tip in [
        "Search is case-insensitive.",
        "Use the media= query parameter to restrict to vinyl or cd.",
        "Matrix runout plain text is searchable via the extra JSONB field.",
        "To find all tokens containing a specific symbol, query token_symbols JOIN symbols WHERE slug = 'emi-triangle'.",
    ]:
        _bullet(doc, tip)

    _h2(doc, "15.3 Adding a New Symbol Manually")
    for step, detail in [
        ("1. Open the symbol registry", "There is currently no UI for symbol management — use the admin shell."),
        ("2. Choose a slug", "Pick a lowercase, hyphenated, descriptive slug (e.g. columbia-pitman).  "
                             "It must be unique and must never change once assigned."),
        ("3. Insert via SQL or admin endpoint",
         "INSERT INTO symbols (slug, name, category, taxonomy_level, region_scope, is_confirmed) "
         "VALUES ('my-slug', 'My Symbol Name', 'pressing_plant_mark', 3, 'US', false);"),
        ("4. Confirm in review",
         "Create a ReviewItem with reason=novel_entity, details pointing to the new symbol.  "
         "A reviewer approves and sets is_confirmed = true."),
    ]:
        _h3(doc, step)
        _body(doc, detail)

    _h2(doc, "15.4 Running the Worker Manually")
    _body(doc, "In development, run the worker in the foreground:")
    _code(doc, "docker compose exec worker python -m mediacat.worker")
    _body(doc, "Or enqueue a single job from the Python shell:")
    _code(doc, "from mediacat.ingestion.queue import enqueue_job, Job")
    _code(doc, "await enqueue_job(Job(connector='discogs', action='fetch_release', payload={'id': 1328315}))")

    _h2(doc, "15.5 Backup & Restore")
    _body(doc,
        "Automated daily backups run in the backup container.  Backups are stored "
        "in the MinIO backups bucket.")
    for cmd, desc in [
        ("make data-init",         "Verify backup directory structure exists"),
        ("deploy/scripts/backup.sh","Run a manual backup (PostgreSQL dump + MinIO snapshot)"),
    ]:
        p = doc.add_paragraph()
        _add_run(p, cmd, mono=True, bold=True, size=10, colour=C_TEAL)
        _add_run(p, f"  —  {desc}", size=10, colour=C_SLATE)
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(3)

    _body(doc,
        "To restore from backup, stop the stack, replace the PostgreSQL data volume "
        "with the backup dump (psql < backup.sql), restart, and run alembic upgrade head "
        "to ensure the schema is current.")

    # ── Footer page ────────────────────────────────────────────────────────
    doc.add_page_break()
    final = doc.add_paragraph()
    final.alignment = WD_ALIGN_PARAGRAPH.CENTER
    final.paragraph_format.space_before = Pt(100)
    _add_run(final, "MediaCat  ·  Technical Reference  ·  v0.1  ·  April 2026",
             size=10, colour=C_SLATE, italic=True)
    doc.add_paragraph()
    notice = doc.add_paragraph()
    notice.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(notice,
        "This document is CONFIDENTIAL.  Do not distribute outside the project team.",
        size=9, colour=C_RED)

    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    out = sys.argv[1] if len(sys.argv) > 1 else "/mnt/c/MEGA/SoundDB/MediaCat_Technical_Reference.docx"
    build(out)
